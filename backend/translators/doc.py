from docx import Document
from docx.shared import RGBColor
from .base import Translator


class DocxTranslator(Translator):
    def translate_docx(self, file_path: str, output_path: str):
        doc = Document(file_path)
        translated_doc = Document()

        # Translate each paragraph
        for paragraph in doc.paragraphs:
            translated_text = self.translate_text(paragraph.text)
            translated_doc.add_paragraph(paragraph.text)
            translated_paragraph = translated_doc.add_paragraph(translated_text)

            # Add color
            for run in translated_paragraph.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)

        # Save the translated document
        translated_doc.save(output_path)
